"""Generate the two binary assets in examples/demo-corpus/ (tempdoc 669).

Manual, throwaway build-time step — NOT run in CI, NOT a tracked dependency of
this repo. It requires packages that are intentionally not added to any
requirements/pyproject file, since the output is a static asset committed once,
not something regenerated on every build:

    pip install python-docx Pillow augraphy

Regenerate only if the demo corpus's fabricated content changes. After running,
recompute examples/demo-corpus/corpus-signature.json.

- harbor-ledger.docx: a born-digital Office document (real python-docx output,
  exercises the Office extraction path with real text, no OCR involved).
- weathered-manifest.png: fabricated cargo-manifest text rendered to an image,
  then degraded via Augraphy (MIT License, https://github.com/sparkfish/augraphy)
  to look like a real aged/scanned document — exercises the real OCR extraction
  path without sourcing (and having to license-clear) an actual scanned
  document. English only: the product's bundled Tesseract currently ships only
  English tessdata (packaging/runtime/tesseract-windows.v1.json).
"""

from __future__ import annotations

from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent.parent.parent / "examples" / "demo-corpus"

MANIFEST_TEXT = """VERRENMOOR CUSTOMS MANIFEST

Vessel: The Halyard Wren
Cargo: 12 crates dried lavender, 4 barrels salt cod, 1 sealed chest (unlisted)
Inspector: C. Aldric
Date: 14 Nov 1907
"""


def generate_docx() -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Verrenmoor Harbor Ledger — Week of 22 April", level=1)
    doc.add_paragraph("Monday: Schooner Wraithwing arrives, unloads tin ore.")
    doc.add_paragraph("Tuesday: Customs delay — inspection of unmarked containers.")
    doc.add_paragraph("Wednesday: Departure of the Halyard Wren, bound for Ostrona.")
    doc.add_paragraph(
        "Notation: Harbor-master Ines Calder recorded a strange humming from "
        "Warehouse Seven all week; cause undetermined."
    )
    out = OUT_DIR / "harbor-ledger.docx"
    doc.save(out)
    print(f"wrote {out}")


def generate_ocr_asset() -> None:
    from PIL import Image, ImageDraw, ImageFont
    from augraphy import AugraphyPipeline, InkBleed, LowInkRandomLines, NoiseTexturize, BrightnessTexturize

    # Render the fabricated manifest text onto a clean white "page".
    width, height = 1200, 900
    img = Image.new("L", (width, height), color=255)
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("cour.ttf", 28)
    except OSError:
        font = ImageFont.load_default()
    draw.multiline_text((60, 60), MANIFEST_TEXT, fill=0, font=font, spacing=14)
    clean = img.convert("RGB")

    import numpy as np

    # A deliberately moderate, seeded pipeline: visibly aged/uneven, not a
    # heavy stamp/ink-blot effect that could make the text unreadable. Fixed
    # random_seed keeps this reproducible run-to-run (the default
    # `default_augraphy_pipeline()` is unseeded and can add heavy overlay
    # elements like ink stamps — deliberately not used here).
    pipeline = AugraphyPipeline(
        ink_phase=[InkBleed(intensity_range=(0.1, 0.2), p=1.0)],
        paper_phase=[
            NoiseTexturize(sigma_range=(3, 6), turbulence_range=(2, 4), p=1.0),
            BrightnessTexturize(texturize_range=(0.9, 0.99), p=1.0),
        ],
        post_phase=[LowInkRandomLines(count_range=(2, 4), p=0.6)],
        random_seed=669,
    )
    degraded = pipeline(np.array(clean))
    result = degraded[0] if isinstance(degraded, (list, tuple)) else degraded

    # JPEG, not PNG: the noise texture makes PNG's lossless compression nearly
    # useless (>1MB), and a real scan/photo of a document is JPEG-shaped
    # anyway — more representative, and keeps the corpus small.
    out = OUT_DIR / "weathered-manifest.jpg"
    Image.fromarray(result).save(out, format="JPEG", quality=82)
    print(f"wrote {out}")


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_docx()
    generate_ocr_asset()
